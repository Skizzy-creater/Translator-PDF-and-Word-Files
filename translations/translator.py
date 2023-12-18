from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import QFileDialog, QMessageBox, QDialog
from PyQt6 import uic
import sys
import docx
from docx import Document
from PyPDF2 import PdfFileReader
from googletrans import Translator




class Ui_Form(object):
    def setupUi(self, Form):
        Form.setObjectName("Form")
        Form.setFixedSize(321, 479)
        font = QtGui.QFont()
        font.setFamily("MS Reference Sans Serif")
        Form.setFont(font)
        self.btn_translateWords = QtWidgets.QPushButton(Form)
        self.btn_translateWords.setEnabled(True)
        self.btn_translateWords.setGeometry(QtCore.QRect(20, 130, 261, 31))
        font = QtGui.QFont()
        font.setFamily("MS Reference Sans Serif")
        font.setPointSize(7)
        self.btn_translateWords.setFont(font)
        self.btn_translateWords.setObjectName("btn_translateWords")
        self.btn_translateWord = QtWidgets.QPushButton(Form)
        self.btn_translateWord.setGeometry(QtCore.QRect(20, 260, 261, 31))
        self.btn_translateWord.setObjectName("btn_translateWord")
        self.btn_translatePDF = QtWidgets.QPushButton(Form)
        self.btn_translatePDF.setGeometry(QtCore.QRect(20, 380, 261, 31))
        self.btn_translatePDF.setObjectName("btn_translatePDF")
        self.progressBar_translateText = QtWidgets.QProgressBar(Form)
        self.progressBar_translateText.setGeometry(QtCore.QRect(20, 100, 301, 23))
        self.progressBar_translateText.setProperty("value", 24)
        self.progressBar_translateText.setObjectName("progressBar_translateText")
        self.progressBar_translateWord = QtWidgets.QProgressBar(Form)
        self.progressBar_translateWord.setGeometry(QtCore.QRect(20, 230, 301, 23))
        self.progressBar_translateWord.setProperty("value", 24)
        self.progressBar_translateWord.setObjectName("progressBar_translateWord")
        self.progressBar_translatePdf = QtWidgets.QProgressBar(Form)
        self.progressBar_translatePdf.setGeometry(QtCore.QRect(20, 350, 301, 23))
        self.progressBar_translatePdf.setProperty("value", 24)
        self.progressBar_translatePdf.setObjectName("progressBar_translatePdf")
        self.label_outpute = QtWidgets.QLabel(Form)
        self.label_outpute.setGeometry(QtCore.QRect(20, 60, 261, 31))
        self.label_outpute.setStyleSheet("background-color: rgb(246, 255, 251);\n"
"font: 75 14pt \"Schadow BT\";")
        self.label_outpute.setObjectName("label_outpute")
        self.lineEdit = QtWidgets.QLineEdit(Form)
        self.lineEdit.setGeometry(QtCore.QRect(20, 20, 261, 31))
        self.lineEdit.setStyleSheet("background-color: rgb(246, 255, 251);\n"
"font: 75 14pt \"Schadow BT\";")
        self.lineEdit.setObjectName("lineEdit")

        self.retranslateUi(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)

        self.add_functions()
        self.progressBar_translatePdf.setValue(0)
        self.progressBar_translateWord.setValue(0)
        self.progressBar_translateText.setValue(0)

    def retranslateUi(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Translator"))
        self.btn_translateWords.setText(_translate("Form", "Перевод"))
        self.btn_translateWord.setText(_translate("Form", "Выберите файл Word"))
        self.btn_translatePDF.setText(_translate("Form", "Выберите файл pdf"))
        self.label_outpute.setText(_translate("Form", "Ответ"))


    def add_functions(self):
        self.btn_translateWords.clicked.connect(self.Translator_text)
        self.btn_translateWord.clicked.connect(self.Browse_Word)
        self.btn_translatePDF.clicked.connect(self.Browse_PDF)


    def Translator_text(self):
        translator = Translator()
        input_word = self.lineEdit.text()
        self.progressBar_translateText.setValue(0)
        if len(input_word) < 23:
            outpute_word = translator.translate(input_word, dest='ru')
            self.label_outpute.setText(outpute_word.text)
            self.progressBar_translateText.setValue(100)
        else:
            dlg = QMessageBox()
            dlg.setWindowTitle("Erorr")
            dlg.setText("Слишком много слов")
            button = dlg.exec()
            self.progressBar_translateText.setValue(0)


    def Browse_Word(self):
        fname = QFileDialog.getOpenFileName()
        translator = Translator()
        if fname:
            document = Document(fname[0])
            paragraphs = document.paragraphs
            document_result = Document()
            paragraphs_resultdocx = document_result.paragraphs
            print('Длина каждого параграфа в тексте', len(paragraphs), '\n')
            i = 0
            while i < len(paragraphs):
                text = paragraphs[i]
                text_translate = translator.translate(text.text, dest='ru')
                document_result.add_paragraph(text_translate.text)
                value = i * 100 / len(paragraphs)
                i += 1
                self.progressBar_translateWord.setValue(int(value))
            else:
                document_result.save('TranslateWord.docx')
                self.progressBar_translateWord.setValue(99)
                dlg = QMessageBox()
                dlg.setWindowTitle("Sucsessful")
                dlg.setText("Файл Word был переведен проверьте папку")
                button = dlg.exec()
                self.progressBar_translateWord.setValue(100)


    def Browse_PDF(self):
        fname = QFileDialog.getOpenFileName()
        translator = Translator()
        if fname:
            with open(fname[0], 'rb') as file:
                pdf = PdfFileReader(file)
                pages = pdf.getNumPages()
                document = docx.Document()
                paragraphs = document.paragraphs
                print("Количество страниц в документе: %i\n\n" % pages)
                for i in range(pages):
                    page = pdf.getPage(i)
                    print("\n\nСтр.", i, "\nПереведенно;\n")
                    text = page.extractText()
                    text_translate = translator.translate(text, dest='ru')
                    document.add_paragraph(text_translate.text)
                    value = i * 100 / pages
                    self.progressBar_translatePdf.setValue(int(value))
                else:
                    document.save('TranslatePDF.docx')
                    self.progressBar_translatePdf.setValue(99)
                    dlg = QMessageBox()
                    dlg.setWindowTitle("Sucsessful")
                    dlg.setText("PDF файл был переведен проверьте папку")
                    button = dlg.exec()
                    self.progressBar_translatePdf.setValue(100)


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_Form()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec())
